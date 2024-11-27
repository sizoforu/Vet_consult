from pydantic import BaseModel, Field, field_validator
from typing import List, Optional, Union
from datetime import date
from enum import Enum
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import PromptTemplate
import streamlit as st
import json
import re
import docx2txt
from PyPDF2 import PdfReader
import os
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import io
import sqlite3
#from datetime import datetime

# Load environment variables
load_dotenv()


class DiagnosticPriority(str, Enum):
    LOW = 'low'
    MEDIUM = 'medium'
    HIGH = 'high'
    URGENT = 'urgent'

#class Diagnostic(BaseModel):
#    diagnosis: str
#    details: str = ''
#    priority: DiagnosticPriority



class Priority(str, Enum):
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    URGENT = "urgent"
    
    @classmethod
    def _missing_(cls, value):
        # Handle case-insensitive matching
        if isinstance(value, str):
            value = value.lower()
            for member in cls:
                if member.value == value:
                    return member
    
class History(BaseModel):
    type: str = Field(..., description="Different Types of history")
    findings: List[str] = Field(..., description="List of findings from the history")
    notes: Optional[Union[str, List[str]]] = Field(None, description="Additional history notes")

class Examination(BaseModel):
    type: str = Field(..., description="Types of examination performed")
    findings: List[str] = Field(..., description="List of findings from the examination")
    region: Optional[str] = Field(None, description="Body region examined")
    notes: Optional[Union[str, List[str]]] = Field(None, description="Additional examination notes")

class Recommendation(BaseModel):
    type: str = Field(..., description="Types of recommendation including medication")
    details: Union[str, List[str]] = Field(..., description="Specific instructions - can be a single string or list of strings")
    duration: Optional[str] = Field(None, description="Time period if applicable")
    notes: Optional[Union[str, List[str]]] = Field(None, description="Additional information")

class Diagnostic(BaseModel):
    type: str = Field(..., description="Types of diagnostic")
    test: str = Field(default="pending", description="Name of the specific test")
    region: Optional[str] = Field(None, description="Body region if applicable")
    priority: Priority = Field(..., description="Urgency level")
    reason: Optional[str] = Field(None, description="Reason for the diagnostic")    


class VeterinaryConsultation(BaseModel):
    # Required fields
    consultation_date: date = Field(..., description="Consultation date in YYYY-MM-DD format")
    veterinarian_name: str = Field(..., description="Veterinarian's name")
    pet_name: str = Field(..., description="Pet's name")
    pet_breed: str = Field(..., description="Pet's breed")
    pet_age: float = Field(..., ge=0, description="Pet's age in years")
    pet_sex: str = Field(..., description="Pet's sex")
    owner_name: str = Field(..., description="Pet owner's name")
    owner_phone: str = Field(..., description="Contact phone number")
    
    # Optional fields with empty lists as defaults
    symptoms: List[str] = Field(default_factory=list, description="List of observed symptoms")
    histories: List[History] = Field(default_factory=list, description="List of histories")
    examinations: List[Examination] = Field(default_factory=list, description="List of examinations performed")
    recommendations: List[Recommendation] = Field(default_factory=list, description="List of recommendations including medications")
    diagnostics: List[Diagnostic] = Field(default_factory=list, description="List of diagnostics")

    @field_validator('owner_phone')
    def validate_phone(cls, v):
        # Remove any non-digit characters except leading '+'
        cleaned_phone = '+' + ''.join(filter(str.isdigit, v)) if v.startswith('+') else ''.join(filter(str.isdigit, v))
        
        # Check if the cleaned number matches the pattern
        if not re.match(r'^\+?1?\d{9,15}$', cleaned_phone):
            raise ValueError('Invalid phone number format')
        return cleaned_phone

    @field_validator('diagnostics')
    def validate_diagnostics(cls, v):
        # Ensure each diagnostic has a test value
        for diagnostic in v:
            if diagnostic.test is None:
                diagnostic.test = "pending"  # Set default value if None
        return v

class VetDatabaseManager:
    def __init__(self, db_name='vet_consult.db'):
        self.db_name = db_name
        self.create_database()

    def create_database(self):
        """Create the veterinary consultation database and table."""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS consultations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            consultation_date DATE NOT NULL,
            veterinarian_name TEXT NOT NULL,
            pet_name TEXT NOT NULL,
            pet_breed TEXT NOT NULL,
            pet_age REAL NOT NULL,
            pet_sex TEXT NOT NULL,
            owner_name TEXT NOT NULL,
            owner_phone TEXT NOT NULL,
            symptoms TEXT NOT NULL,
            histories TEXT NOT NULL,
            examinations TEXT NOT NULL,
            recommendations TEXT NOT NULL,
            diagnostics TEXT NOT NULL
        )
        ''')
        
        conn.commit()
        conn.close()
        
    def _format_history(self, hist):
        """Format history data in a human-readable format."""
        parts = [f"Type: {hist.type}"]
        if hist.findings:
            parts.append("Findings:")
            parts.extend([f"  - {finding}" for finding in hist.findings])
        if hist.notes:
            parts.append(f"Notes: {hist.notes}")
        return "\n".join(parts)
 
    
    def _format_examination(self, exam):
        """Format examination data in a human-readable format."""
        parts = [f"Type: {exam.type}"]
        if exam.findings:
            parts.append("Findings:")
            parts.extend([f"  - {finding}" for finding in exam.findings])
        if exam.region:
            parts.append(f"Region: {exam.region}")
        if exam.notes:
            parts.append(f"Notes: {exam.notes}")
        return "\n".join(parts)

    def _format_recommendation(self, rec):
        """Format recommendation data in a human-readable format."""
        parts = [f"Type: {rec.type}"]
        parts.append(f"Details: {rec.details}")
        if rec.duration:
            parts.append(f"Duration: {rec.duration}")
        if rec.notes:
            parts.append(f"Notes: {rec.notes}")
        return "\n".join(parts)

    def _format_diagnostic(self, diag):
        """Format diagnostic data in a human-readable format."""
        parts = [f"Type: {diag.type}"]
        parts.append(f"Test: {diag.test}")
        parts.append(f"Priority: {diag.priority}")
        parts.append(f"Reason: {diag.reason}")
        if diag.region:
            parts.append(f"Region: {diag.region}")
        return "\n".join(parts)
    
    
    def _parse_history(self, text):
        """Parse history text back into a structured format."""
        lines = text.split('\n')
        hist_dict = {}
        findings = []
        current_section = None
        
        for line in lines:
            if line.startswith('Type: '):
                hist_dict['type'] = line.replace('Type: ', '')
            elif line.startswith('Notes: '):
                hist_dict['notes'] = line.replace('Notes: ', '')
            elif line == 'Findings:':
                current_section = 'findings'
            elif line.startswith('  - ') and current_section == 'findings':
                findings.append(line.replace('  - ', ''))
        
            hist_dict['findings'] = findings
        return History(**hist_dict)
    

    def _parse_examination(self, text):
        """Parse examination text back into a structured format."""
        lines = text.split('\n')
        exam_dict = {}
        findings = []
        current_section = None
        
        for line in lines:
            if line.startswith('Type: '):
                exam_dict['type'] = line.replace('Type: ', '')
            elif line.startswith('Region: '):
                exam_dict['region'] = line.replace('Region: ', '')
            elif line.startswith('Notes: '):
                exam_dict['notes'] = line.replace('Notes: ', '')
            elif line == 'Findings:':
                current_section = 'findings'
            elif line.startswith('  - ') and current_section == 'findings':
                findings.append(line.replace('  - ', ''))
        
        exam_dict['findings'] = findings
        return Examination(**exam_dict)



    def _parse_recommendation(self, text):
        """Parse recommendation text back into a structured format."""
        # If no text, return an empty Recommendation
        if not text:
            return Recommendation(type='', details='No details provided')

        lines = text.split('\n')
        rec_dict = {}
        
        for line in lines:
            if line.startswith('Type: '):
                rec_dict['type'] = line.replace('Type: ', '').strip()
            elif line.startswith('Details: '):
                rec_dict['details'] = line.replace('Details: ', '').strip()
            elif line.startswith('Duration: '):
                rec_dict['duration'] = line.replace('Duration: ', '').strip()
            elif line.startswith('Notes: '):
                rec_dict['notes'] = line.replace('Notes:', '').strip()
            
        # Ensure required fields are present
        if 'type' not in rec_dict:
             rec_dict['type'] = ''
             
             
        # Provide a default details if not present
        if 'details' not in rec_dict:
            if rec_dict['type'] == "uso de collar isabelino":
                rec_dict['details'] = "usar continuamente hasta mejoría de lesiones"
            elif rec_dict['type'] == "suplemento de ácidos grasos omega 3 y 6":
                rec_dict['details'] = "administrar según indicaciones del fabricante"
            else:
                rec_dict['details'] = "detalles pendientes de especificar"
                
                
        # Return statement should be outside the for loop        
        return Recommendation(**rec_dict)
            
        
    def _parse_diagnostic(self, diag_text):
        """Parse diagnostic text into a Diagnostic object."""
       # If no text, return None instead of empty Diagnostic since we have required fields
        if not diag_text:
            return None

        lines = diag_text.split('\n')
        diag_dict = {}
        
        for line in lines:
            if line.startswith('Type: '):
                diag_dict['type'] = line.replace('Type: ', '').strip()
            elif line.startswith('Test: '):
                diag_dict['test'] = line.replace('Test: ', '').strip()
            elif line.startswith('Priority: '):
                priority_value = line.replace('Priority: ', '').strip().lower()
                # Convert Priority.LOW to 'low' if needed
                if priority_value.startswith('priority.'):
                    priority_value = priority_value.split('.')[-1].lower()
                
                # Only set priority if it's a valid enum value
                if priority_value in ['low', 'medium', 'high', 'urgent']:
                    diag_dict['priority'] = priority_value
            elif line.startswith('Reason: '):
                diag_dict['reason'] = line.replace('Reason: ', '').strip()
            elif line.startswith('Region: '):
                diag_dict['region'] = line.replace('Region: ', '').strip()
              
                 
        # No need to raise ValueError for missing fields since they're optional now      
        return Diagnostic(**diag_dict)

    
    def add_consultation(self, consultation: VeterinaryConsultation):
        """Add a new consultation to the database with human-readable format."""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        # Format complex fields in human-readable format
        formatted_symptoms = "\n- " + "\n- ".join(consultation.symptoms)
        formatted_examinations = "\n\n".join(self._format_examination(exam) for exam in consultation.examinations)
        formatted_histories = "\n\n".join(self._format_history(hist) for hist in consultation.histories)  
        formatted_recommendations = "\n\n".join(self._format_recommendation(rec) for rec in consultation.recommendations)
        formatted_diagnostics = "\n\n".join(self._format_diagnostic(diag) for diag in consultation.diagnostics)
        
        cursor.execute('''
        INSERT INTO consultations (
            consultation_date,
            veterinarian_name,
            pet_name,
            pet_breed,
            pet_age,
            pet_sex,
            owner_name,
            owner_phone,
            symptoms,
            histories,
            examinations,
            recommendations,
            diagnostics
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            consultation.consultation_date.isoformat(),
            consultation.veterinarian_name,
            consultation.pet_name,
            consultation.pet_breed,
            consultation.pet_age,
            consultation.pet_sex,
            consultation.owner_name,
            consultation.owner_phone,
            formatted_symptoms,
            formatted_histories,
            formatted_examinations,
            formatted_recommendations,
            formatted_diagnostics
        ))
        
        consultation_id = cursor.lastrowid
        conn.commit()
        conn.close()
        return consultation_id
    
    
    
    def update_consultation(self, consultation_id: int, consultation: VeterinaryConsultation):
        """Update an existing consultation in the database with human-readable format."""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        try:
            # Format complex fields in human-readable format
            formatted_symptoms = "\n- " + "\n- ".join(consultation.symptoms) if consultation.symptoms else ""
            
            # Handle potentially empty lists
            formatted_histories = "\n\n".join(
                self._format_history(hist) for hist in (consultation.histories or [])
            )
            
            # Handle potentially empty lists
            formatted_examinations = "\n\n".join(
                self._format_examination(exam) for exam in (consultation.examinations or [])
            )
            
            formatted_recommendations = "\n\n".join(
                self._format_recommendation(rec) for rec in (consultation.recommendations or [])
            )
            
            formatted_diagnostics = "\n\n".join(
                self._format_diagnostic(diag) for diag in (consultation.diagnostics or [])
            )
            
            cursor.execute('''
            UPDATE consultations SET
                consultation_date = ?,
                veterinarian_name = ?,
                pet_name = ?,
                pet_breed = ?,
                pet_age = ?,
                pet_sex = ?
                owner_name = ?,
                owner_phone = ?,
                symptoms = ?,
                histories = ?
                examinations = ?,
                recommendations = ?,
                diagnostics = ?
            WHERE id = ?
            ''', (
                consultation.consultation_date.isoformat(),
                consultation.veterinarian_name,
                consultation.pet_name,
                consultation.pet_breed,
                consultation.pet_age,
                consultation.pet_sex,
                consultation.owner_name,
                consultation.owner_phone,
                formatted_symptoms,
                formatted_histories,
                formatted_examinations,
                formatted_recommendations,
                formatted_diagnostics,
                consultation_id
            ))
            
            rows_affected = cursor.rowcount
            conn.commit()
            return rows_affected > 0
            
        except Exception as e:
            conn.rollback()
            raise e
        finally:
            conn.close()
    
    def get_consultation(self, consultation_id: int) -> VeterinaryConsultation:
        """Retrieve a specific consultation from the database."""
        conn = sqlite3.connect(self.db_name)
        try:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM consultations WHERE id = ?', (consultation_id,))
            row = cursor.fetchone()
    
            if row:
                columns = [description[0] for description in cursor.description]
                consultation_dict = dict(zip(columns, row))
        
            # Parse symptoms from human-readable format
            if consultation_dict['symptoms']:
                consultation_dict['symptoms'] = [
                    s.strip('- ') for s in consultation_dict['symptoms'].split('\n') 
                    if s.strip('- ')
                ]
        
            # Parse histories from human-readable format
            if consultation_dict['histories']:
                consultation_dict['histories'] = [
                    self._parse_history(hist_text)
                    for hist_text in consultation_dict['histories'].split('\n\n')
                    if hist_text.strip()
                ]
        
            # Parse examinations from human-readable format
            if consultation_dict['examinations']:
                consultation_dict['examinations'] = [
                    self._parse_examination(exam_text)
                    for exam_text in consultation_dict['examinations'].split('\n\n')
                    if exam_text.strip()
                ]
        
            # Parse recommendations from human-readable format
            if consultation_dict['recommendations']:
                consultation_dict['recommendations'] = [
                    self._parse_recommendation(rec_text)
                    for rec_text in consultation_dict['recommendations'].split('\n\n')
                    if rec_text.strip()
                ]
        
            # Parse diagnostics from human-readable format
            if consultation_dict['diagnostics']:
                consultation_dict['diagnostics'] = [
                    self._parse_diagnostic(diag_text)
                    for diag_text in consultation_dict['diagnostics'].split('\n\n')
                    if diag_text.strip()
                ]            
            
            # Convert date string back to date object
            consultation_dict['consultation_date'] = date.fromisoformat(consultation_dict['consultation_date'])
        
            return VeterinaryConsultation(**consultation_dict)
            return None
        finally:
            conn.close()
    
    
    
    
    def get_all_consultations_summary(self):
        """Get a summary of all consultations for display in a selection widget."""
        conn = sqlite3.connect(self.db_name)
        try:
            cursor = conn.cursor()
            cursor.execute('''
            SELECT id, consultation_date, pet_name, owner_name
            FROM consultations
            ORDER BY consultation_date DESC
            ''')
            consultations = cursor.fetchall()
            return [{"id": c[0], "date": c[1], "pet": c[2], "owner": c[3]} for c in consultations]
        except Exception as e:
            print(f"Error retrieving consultations summary: {e}")
            return []
        finally:
            conn.close()  
            
              

    def get_all_consultations(self) -> List[VeterinaryConsultation]:
        """Retrieve all consultations from the database."""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM consultations')
        rows = cursor.fetchall()
        
        consultations = []
        columns = [description[0] for description in cursor.description]
        
        for row in rows:
            consultation_dict = dict(zip(columns, row))
            
            # Parse JSON strings back to Python objects
            consultation_dict['symptoms'] = json.loads(consultation_dict['symptoms'])
            consultation_dict['histories'] = json.loads(consultation_dict['histories'])
            consultation_dict['examinations'] = json.loads(consultation_dict['examinations'])
            consultation_dict['recommendations'] = json.loads(consultation_dict['recommendations'])
            consultation_dict['diagnostics'] = json.loads(consultation_dict['diagnostics'])
            
            # Convert date string back to date object
            consultation_dict['consultation_date'] = date.fromisoformat(consultation_dict['consultation_date'])
            
            consultations.append(VeterinaryConsultation(**consultation_dict))
        
        conn.close()
        return consultations


def display_consultation_data(data):
    """Display the extracted consultation data in a structured format."""
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.subheader("📌 Informacion Basica")
        st.write(f"📅 Fecha: {data.consultation_date}")
        st.write(f"👨‍⚕️ Veterinario: {data.veterinarian_name}")
        
        st.subheader("👤 Informacion Duena")
        st.write(f"Nombre: {data.owner_name}")
        st.write(f"Telefono: {data.owner_phone}")

    with col2:
        st.subheader("🐾 Informacion Mascota")
        st.write(f"Nombre: {data.pet_name}")
        st.write(f"Raza: {data.pet_breed}")
        st.write(f"Edad: {data.pet_age} years")
        st.write(f"Sexo: {data.pet_sex}")
        
        st.subheader("🤒 Sintomas")
        for symptom in data.symptoms:
            st.write(f"• {symptom}")


        st.subheader("🔍 Historias")
        for hist in data.histories:
            with st.expander(f"{hist.type}"):
                st.write("resultados:")
                for finding in hist.findings:
                    st.write(f"• {finding}")
                if hist.notes:
                    st.write(f"Notas: {hist.notes}")


    with col3:
        st.subheader("🔍 Exámenes")
        for exam in data.examinations:
            with st.expander(f"{exam.type}"):
                st.write("Resultados:")
                for finding in exam.findings:
                    st.write(f"• {finding}")
                if exam.region:
                    st.write(f"Región: {exam.region}")
                if exam.notes:
                    st.write(f"Notas: {exam.notes}")

        st.subheader("💡 Recomendaciones")
        for rec in data.recommendations:
            with st.expander(f"{rec.type}"):
                st.write(f"Detalles: {rec.details}")
                if rec.duration:
                    st.write(f"Duración: {rec.duration}")
                if rec.notes:
                    st.write(f"Notas: {rec.notes}")

        st.subheader("🔬 Diagnosticos")
        for diag in data.diagnostics:
            with st.expander(f"{diag.type}"):
                st.write(f"Testeo: {diag.test}")
                st.write(f"Prioridad: {diag.priority}")
                st.write(f"Razón: {diag.reason}")
                if diag.region:
                    st.write(f"Región: {diag.region}")

class VeterinaryConsultationExtractor:
    def __init__(self):
        self.parser = PydanticOutputParser(pydantic_object=VeterinaryConsultation)
        self.prompt_template = self._create_prompt_template()
        self.llm = self._initialize_llm()

    def _create_prompt_template(self):
        template = """
        Extract strucure data from veterinary consultations from the following text and format it according to these requirements:

        {format_instructions}

        Important guidelines:
        1. Extract exact dates, names, and numbers when present
        2. Categorize symptoms clearly and concisely
        3. Categorize types of histories and their findings accordingly to the text
        4. Categorize types of examinations and their findings accordingly to the text
        5. Categorize recommendations with medications and clear timeframes when specified
        6. Assign appropriate priority levels to diagnostics

        Text: {consultation_text}

        Please provide a complete and detailed extraction following the exact schema specified.
        Use the current date if no specific date is mentioned.
        Do not make assumptions for required fields if the information is not explicitly stated.
        """
        return PromptTemplate(
            template=template,
            input_variables=["consultation_text"],
            partial_variables={"format_instructions": self.parser.get_format_instructions()}
        )

    def _initialize_llm(self):
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY not found in environment variables")
        
        return ChatOpenAI(
            api_key=api_key,
            base_url="https://api.groq.com/openai/v1",
            model="llama3-8b-8192",
            #model="mixtral-8x7b-32768",
            #model="llama3-70b-8192",
            #model="gemma2-9b-it",
            temperature=0.1
        )

    def extract_text_from_file(self, file):
        """Extract text from uploaded PDF or DOCX file."""
        try:
            if file.type == "application/pdf":
                pdf_reader = PdfReader(file)
                return ' '.join(page.extract_text() for page in pdf_reader.pages)
            else:  # DOCX
                return docx2txt.process(file)
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")

    def process_consultation(self, text):
        """Process consultation text and extract structured information."""
        try:
            # Generate formatted prompt
            formatted_prompt = self.prompt_template.format_prompt(
                consultation_text=text
            ).to_string()
            
            # Get response from LLM
            response = self.llm.invoke(formatted_prompt)
            
            try:
                # Try to parse JSON from the response
                json_str = re.search(r'\{.*\}', response.content, re.DOTALL)
                if json_str:
                    consultation_data = json.loads(json_str.group())
                    # Parse the cleaned JSON into the Pydantic model
                    return self.parser.parse(json.dumps(consultation_data))
                else:
                    raise ValueError("No JSON found in response")
            except json.JSONDecodeError as e:
                st.error(f"JSON parsing error: {str(e)}")
                return None
            
        except Exception as e:
            st.error(f"Error processing consultation: {str(e)}")
            return None

def create_consultation_docx(consultation_data):
        """
        Create a formatted DOCX document from consultation data.
        Returns a BytesIO object containing the document.
        """
        doc = Document()
    
        # Set up title
        title = doc.add_heading('Veterinary Consultation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        # Add basic information section
        doc.add_heading('Basic Information', level=1)
        basic_info = doc.add_paragraph()
        basic_info.add_run('Fecha: ').bold = True
        basic_info.add_run(f"{consultation_data.consultation_date}\n")
        basic_info.add_run('Veterinario: ').bold = True
        basic_info.add_run(f"{consultation_data.veterinarian_name}\n")
    
        # Add owner information
        doc.add_heading('Information Duena', level=1)
        owner_info = doc.add_paragraph()
        owner_info.add_run('Nombre: ').bold = True
        owner_info.add_run(f"{consultation_data.owner_name}\n")
        owner_info.add_run('Telefono: ').bold = True
        owner_info.add_run(f"{consultation_data.owner_phone}\n")
    
        # Add pet information
        doc.add_heading('Information Mascota', level=1)
        pet_info = doc.add_paragraph()
        pet_info.add_run('Nombre: ').bold = True
        pet_info.add_run(f"{consultation_data.pet_name}\n")
        pet_info.add_run('Raza: ').bold = True
        pet_info.add_run(f"{consultation_data.pet_breed}\n")
        pet_info.add_run('Edad: ').bold = True
        pet_info.add_run(f"{consultation_data.pet_age} years\n")
    
        # Add symptoms
        doc.add_heading('Symptomas', level=1)
        for symptom in consultation_data.symptoms:
            doc.add_paragraph(symptom, style='List Bullet')
    
    
        # Add histories
        doc.add_heading('Historias', level=1)
        for hist in consultation_data.histories:
            p = doc.add_paragraph()
            p.add_run(f"Tipo: {hist.type}\n").bold = True
            p.add_run('Resultados:\n')
        for finding in hist.findings:
            doc.add_paragraph(finding, style='List Bullet')
        if hist.notes:
            p.add_run(f"Notas: {hist.notes}\n")
        doc.add_paragraph()  # Add spacing between histories
    
    
        # Add examinations
        doc.add_heading('Examinaciones', level=1)
        for exam in consultation_data.examinations:
            p = doc.add_paragraph()
            p.add_run(f"Tipo: {exam.type}\n").bold = True
            p.add_run('Resultados:\n')
        for finding in exam.findings:
            doc.add_paragraph(finding, style='List Bullet')
        if exam.region:
            p.add_run(f"Región: {exam.region}\n")
        if exam.notes:
            p.add_run(f"Notas: {exam.notes}\n")
        doc.add_paragraph()  # Add spacing between examinations
    
    
        # Add recommendations
        doc.add_heading('Recommendations', level=1)
        for rec in consultation_data.recommendations:
            p = doc.add_paragraph()
            p.add_run(f"Tpo: {rec.type}\n").bold = True
            p.add_run(f"Detalles: {rec.details}\n")
        if rec.duration:
            p.add_run(f"Duración: {rec.duration}\n")
        if rec.notes:
            p.add_run(f"Notas: {rec.notes}\n")
        doc.add_paragraph()  # Add spacing between recommendations
    
    
        # Add diagnostics
        doc.add_heading('Diagnosticos', level=1)
        for diag in consultation_data.diagnostics:
            p = doc.add_paragraph()
            p.add_run(f"Tipo: {diag.type}\n").bold = True
            p.add_run(f"Testeo: {diag.test}\n")
            # Add priority with color coding
            priority_run = p.add_run(f"Priority: {diag.priority}\n")
        if diag.priority == "urgent":
            priority_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        elif diag.priority == "high":
            priority_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        p.add_run(f"Razon: {diag.reason}\n")
        if diag.region:
            p.add_run(f"Región: {diag.region}\n")
        doc.add_paragraph()  # Add spacing between diagnostics
    
        # Save to BytesIO object
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
    
        return docx_file


def create_edit_form(consultation_data):
    """
    Create a form for editing the consultation data.
    Returns the updated consultation data or None if the form is not submitted.
    """
    with st.form("edit_consultation_form"):
        st.subheader("Edit Consultation")

        if consultation_data is None:
            st.info("No consultation data available to edit.")
            return None

        # Basic information
        col1, col2 = st.columns(2)
        
        with col1:
            consultation_date = st.date_input("Fecha Consultacion ", value=consultation_data.consultation_date)
            veterinarian_name = st.text_input("Nombre Veterinario", value=consultation_data.veterinarian_name)
        
            # Owner information
            st.subheader("Information Duena")
            owner_name = st.text_input("Nombre Duena", value=consultation_data.owner_name)
            owner_phone = st.text_input("Telefono Duena", value=consultation_data.owner_phone)

        with col2:
            # Pet information
            st.subheader(" Information Mascota")
            pet_name = st.text_input("Mascota Nombre", value=consultation_data.pet_name)
            pet_breed = st.text_input("Mascota Raza", value=consultation_data.pet_breed)
            pet_age = st.number_input("Mascota Edad", value=consultation_data.pet_age, min_value=0.0, step=0.1)
            pet_sex = st.text_input("Mascota Sexo", value=consultation_data.pet_sex)

                  
        # Symptoms
        st.subheader("Sintomas")
        symptoms = st.text_area("Symptomas (comma-separated)", value=", ".join(consultation_data.symptoms))

        
        # Histories
        st.subheader("Historias")
        histories = []
        for i, hist in enumerate(consultation_data.histories):
            with st.expander(f"Historias {i+1}", expanded=True):
                hist_type = st.text_input("Tipo", value=hist.type, key=f"hist_type_{i}")
                findings = st.text_area("Resultados (comma-separated)", 
                                      value=", ".join(hist.findings), 
                                      key=f"hist_findings_{i}")
                notes = st.text_area("Notas", value=hist.notes or "", key=f"hist_notes_{i}")
                histories.append({
                    'type': hist_type,
                    'findings': findings,
                    'notes': notes
                })       
        
        
        # Examinations
        st.subheader("Examinaciones")
        examinations = []
        for i, exam in enumerate(consultation_data.examinations):
            with st.expander(f"Examinacion {i+1}", expanded=True):
                exam_type = st.text_input("Tipo", value=exam.type, key=f"exam_type_{i}")
                findings = st.text_area("Resultados (comma-separated)", 
                                      value=", ".join(exam.findings), 
                                      key=f"exam_findings_{i}")
                region = st.text_input("Region", value=exam.region or "", key=f"exam_region_{i}")
                notes = st.text_area("Notas", value=exam.notes or "", key=f"exam_notes_{i}")
                examinations.append({
                    'type': exam_type,
                    'findings': findings,
                    'region': region,
                    'notes': notes
                })


        # Recommendations
        st.subheader("Recommendaciones")
        recommendations = []
        for i, rec in enumerate(consultation_data.recommendations):
            with st.expander(f"Recommendacion {i+1}", expanded=True):
                rec_type = st.text_input("Tipo", value=rec.type, key=f"rec_type_{i}")
                details = st.text_area("Detalles", value=rec.details, key=f"rec_details_{i}")
                duration = st.text_input("Duracion", value=rec.duration or "", key=f"rec_duration_{i}")
                notes = st.text_area("Notas", value=rec.notes or "", key=f"rec_notes_{i}")
                recommendations.append({
                    'type': rec_type,
                    'details': details,
                    'duration': duration,
                    'notes': notes
                })

        # Diagnostics
        st.subheader("Diagnosticos")
        diagnostics = []
        for i, diag in enumerate(consultation_data.diagnostics):
            with st.expander(f"Diagnostico {i+1}", expanded=True):
                diag_type = st.text_input("Tipo", value=diag.type, key=f"diag_type_{i}")
                test = st.text_input("Testeo", value=diag.test, key=f"diag_test_{i}")
                priority = st.selectbox(
                    "Priority",
                    options=[p.value for p in Priority],
                    index=[p.value for p in Priority].index(diag.priority),
                    key=f"diag_priority_{i}"
                )
                reason = st.text_area("Razon", value=diag.reason, key=f"diag_reason_{i}")
                region = st.text_input("Region", value=diag.region or "", key=f"diag_region_{i}")
                diagnostics.append({
                    'type': diag_type,
                    'test': test,
                    'priority': priority,
                    'reason': reason,
                    'region': region
                })

        submit_button = st.form_submit_button("Guardiar Cambios")
        
        if submit_button:
            try:
                # Convert the form data back to a VeterinaryConsultation object
                return VeterinaryConsultation(
                    consultation_date=consultation_date,
                    veterinarian_name=veterinarian_name,
                    pet_name=pet_name,
                    pet_breed=pet_breed,
                    pet_age=pet_age,
                    pet_sex=pet_sex,
                    owner_name=owner_name,
                    owner_phone=owner_phone,
                    symptoms=[s.strip() for s in symptoms.split(",") if s.strip()],
                    
                    histories=[
                        History(
                            type=e['type'],
                            findings=[f.strip() for f in e['findings'].split(",") if f.strip()],
                            notes=e['notes'] if e['notes'] else None
                        ) for e in histories
                    ],
                    
                    examinations=[
                        Examination(
                            type=e['type'],
                            findings=[f.strip() for f in e['findings'].split(",") if f.strip()],
                            region=e['region'] if e['region'] else None,
                            notes=e['notes'] if e['notes'] else None
                        ) for e in examinations
                    ],
                    
                    recommendations=[
                        Recommendation(
                            type=r['type'],
                            details=r['details'],
                            duration=r['duration'] if r['duration'] else None,
                            notes=r['notes'] if r['notes'] else None
                        ) for r in recommendations
                    ],
                    diagnostics=[
                        Diagnostic(
                            type=d['type'],
                            test=d['test'],
                            priority=Priority(d['priority']),
                            reason=d['reason'],
                            region=d['region'] if d['region'] else None
                        ) for d in diagnostics
                    ]
                )
            except Exception as e:
                st.error(f"Error saving changes: {str(e)}")
                return None
        return None
    pass

def process_uploaded_file():
    """Function to handle file upload and processing"""
    # File upload
    uploaded_file = st.file_uploader("Elige un archivo", type=["pdf", "docx"])
    
    if uploaded_file is not None:
        # Extract text from the uploaded file
        extractor = VeterinaryConsultationExtractor()
        try:
            with st.spinner("Processing documento..."):
                consultation_text = extractor.extract_text_from_file(uploaded_file)
                consultation_data = extractor.process_consultation(consultation_text)
    
            if consultation_data:
                # Display the consultation data
                display_consultation_data(consultation_data)

                # Initialize database manager and save consultation
                db_manager = VetDatabaseManager()
                try:
                    consultation_id = db_manager.add_consultation(consultation_data)
                    st.session_state.consultation_id = consultation_id
                    st.success("✅ Consulta guardada en la base de datos!")
                except Exception as e:
                    st.error(f"Error al guardar en la base de datos: {str(e)}")

                # Create DOCX file
                docx_file = create_consultation_docx(consultation_data)

                # Add download button for the report
                st.download_button(
                    label="📥 Download DOCX Report",
                    data=docx_file.getvalue(),
                    file_name=f"consultation_{consultation_data.consultation_date}_{consultation_data.pet_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            st.exception(e)
    
def edit_consultation():
    """Function to handle consultation editing"""
    # Initialize session state variables
    if 'consultation_selector' not in st.session_state:
        st.session_state.consultation_selector = None
    if 'previous_consultation_id' not in st.session_state:
        st.session_state.previous_consultation_id = None
    if 'consultation_data' not in st.session_state:
        st.session_state.consultation_data = None
    if 'consultations' not in st.session_state:
        st.session_state.consultations = None
    if 'consultation_id' not in st.session_state:
        st.session_state.consultation_id = None

    # Initialize database manager
    db_manager = VetDatabaseManager()
    
    # Fetch all consultations and store in session state
    consultations = db_manager.get_all_consultations_summary()
    st.session_state.consultations = consultations

    def on_consultation_select():
        """Callback function for when a new consultation is selected"""
        selected = st.session_state.consultation_selector
        if selected and selected['id'] != st.session_state.previous_consultation_id:
            consultation_data = db_manager.get_consultation(selected['id'])
            st.session_state.consultation_data = consultation_data
            st.session_state.consultation_id = selected['id']
            st.session_state.previous_consultation_id = selected['id']

    # Create selection box for consultations with callback
    st.selectbox(
        "Select Consultation to Edit",
        options=st.session_state.consultations,
        format_func=lambda x: f"{x['date']} - {x['pet']} ({x['owner']})",
        key="consultation_selector",
        on_change=on_consultation_select
    )

    if st.session_state.consultation_data:
        # Create edit form with current consultation data
        updated_data = create_edit_form(st.session_state.consultation_data)

        if updated_data:
            try:
                success = db_manager.update_consultation(
                    st.session_state.consultation_id,
                    updated_data
                )
                if success:
                    st.success("✅ Los cambios se guardaron correctamente!")
                    st.session_state.consultation_data = updated_data

                    # Create updated DOCX file
                    docx_file = create_consultation_docx(updated_data)

                    # Add download button for updated report
                    st.download_button(
                        label="📥 Download Updated DOCX Report",
                        data=docx_file.getvalue(),
                        file_name=f"consultation_{updated_data.consultation_date}_{updated_data.pet_name}_updated.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("No se actualizó ninguna consulta. Por favor verifique el ID de la consulta.")
            except Exception as e:
                st.error(f"Error saving changes: {str(e)}")

def main():
    #st.set_page_config(page_title="Veterinary Consultation Scanner", layout="wide")

    # Initialize session state
    if 'consultation_data' not in st.session_state:
        st.session_state.consultation_data = None
    if 'consultation_id' not in st.session_state:
        st.session_state.consultation_id = None

    # Create tabs
    scan_tab, edit_tab = st.tabs(["📄 Scan Document", "✏️ Edit Consultation"])

    with scan_tab:
        st.title("📄 Scan Veterinaria Consultacion")
        process_uploaded_file()

    with edit_tab:
        st.title("✏️ Edit Consultacion Datos")
        edit_consultation()

if __name__ == "__main__":
    main()
