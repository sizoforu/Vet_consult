import streamlit as st
import os
import base64
#from pydub import AudioSegment
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import tempfile
#from reportlab.pdfgen import canvas
#from reportlab.lib.pagesizes import letter
#from reportlab.lib import colors
#from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
#from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import extraction

import editing

# Load environment variables
load_dotenv()


# Configuración de la aplicación - ESTE DEBE SER EL PRIMER COMANDO DE STREAMLIT
st.set_page_config(
    layout="wide",
    page_title="Aplicación Médica Veterinaria",
    page_icon="🐾"
)

class AplicacionTranscripcionVeterinaria:
    def __init__(self):
        self.groq_client = OpenAI(
            api_key=st.secrets["GROQ_API_KEY"],
            base_url="https://api.groq.com/openai/v1"
        )

    @staticmethod
    def generar_docx(transcripcion, analisis=None):
        """Genera un documento DOCX a partir de la transcripción y el análisis."""
        doc = Document()
        doc.add_heading('Informe de Transcripción Veterinaria', 0)
        
        doc.add_heading('Transcripción Original', level=1)
        doc.add_paragraph(transcripcion)
        
        if analisis:
            doc.add_heading('Análisis', level=1)
            doc.add_paragraph(analisis)
        
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio

    @staticmethod
    def audio_a_base64(ruta_archivo):
        """Convierte archivo de audio a cadena base64."""
        with open(ruta_archivo, "rb") as archivo_audio:
            return base64.b64encode(archivo_audio.read()).decode()

    def transcribir_audio(self, ruta_archivo_audio):
        """Transcribe audio usando la API de Groq."""
        try:
            with open(ruta_archivo_audio, "rb") as archivo_audio:
                transcripcion = self.groq_client.audio.transcriptions.create(
                    model="whisper-large-v3",
                    file=archivo_audio,
                    response_format="text"
                )
            return transcripcion
        except Exception as e:
            st.error(f"Error durante la transcripción: {str(e)}")
            return None

def mostrar_pagina_transcripcion():
    """Muestra el contenido de la página de transcripción"""
    st.title("🎙️ Transcripción de Audio")
    st.markdown("""
        Sube una grabación de audio de tu consulta veterinaria para generar un informe de transcripción.
    """)
    
    app = AplicacionTranscripcionVeterinaria()
    
    archivo_subido = st.file_uploader("Sube un archivo MP3", type=["mp3"])
    col1, col2 = st.columns(2)

    if archivo_subido is not None:
        with col1:
            st.subheader("📂 Archivo de Audio Subido")
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as archivo_temporal:
                archivo_temporal.write(archivo_subido.getvalue())
                ruta_archivo_temporal = archivo_temporal.name

            audio_base64 = app.audio_a_base64(ruta_archivo_temporal)
            audio_html = f"""
                <audio controls style="width: 100%">
                    <source src="data:audio/mp3;base64,{audio_base64}" type="audio/mp3">
                    Tu navegador no soporta el elemento de audio.
                </audio>
            """
            st.markdown(audio_html, unsafe_allow_html=True)

            if st.button("🎯 Generar Transcripción"):
                with st.spinner("Transcribiendo audio..."):
                    transcripcion = app.transcribir_audio(ruta_archivo_temporal)
                    
                    if transcripcion:
                        with col2:
                            st.subheader("📝 Resultado de la Transcripción")
                            st.success("¡Transcripción completada con éxito!")
                            st.text_area("Transcripción Original", transcripcion, height=300)

                            # Generar archivo DOCX
                            doc_bio = app.generar_docx(transcripcion)
                            
                            # Agregar botón de descarga para DOCX
                            st.download_button(
                                label="💾 Descargar como DOCX",
                                data=doc_bio,
                                file_name="informe_veterinario.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

        os.unlink(ruta_archivo_temporal)

def main():
    # Crear un estilo CSS personalizado para el menú desplegable
    st.markdown("""
        <style>
        .stSelectbox {
            width: 200px;
        }
        </style>
    """, unsafe_allow_html=True)

    # Agregar logo/título en la barra lateral
    st.sidebar.title("🐾 Aplicación Médica Veterinaria")
    
    # Actualizar menú desplegable con la nueva opción
    pagina = st.sidebar.selectbox(
        "Elegir un Módulo",
        ["Audio Transcription", "Document Scanner", "Consultation Editor"],
        format_func=lambda x: {
            "Audio Transcription": "📝 Transcripción de Audio",
            "Document Scanner": "🏥 Escáner de Documentos",
            "Consultation Editor": "✏️ Editor de Consultas"
        }[x]
    )

    if pagina == "Audio Transcription":
        mostrar_pagina_transcripcion()
    elif pagina == "Document Scanner":
        extraction.main()
    else:
        editing.main()

if __name__ == "__main__":
    main()
