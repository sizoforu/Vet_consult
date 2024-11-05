import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
import docx2txt
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage
from langchain_core.prompts import PromptTemplate
import json
import os
import re
from datetime import datetime
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import sqlite3  # Import SQLite library

# Set page config first
st.set_page_config(layout="wide")

# Load environment variables
load_dotenv()

def extract_text_from_file(file):
    """Extract text from uploaded PDF or DOCX file."""
    text = ""
    try:
        if file.type == "application/pdf":
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        else:
            text = docx2txt.process(file)
    except Exception as e:
        st.error(f"❌ Error extracting text from file: {str(e)}")
    return text

def initialize_llm():
    """Initialize the LLM with proper configuration."""
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        st.error("❌ GROQ_API_KEY not found in environment variables")
        return None
    
    try:
        return ChatOpenAI(
            api_key=api_key,
            base_url="https://api.groq.com/openai/v1",
            model="llama3-8b-8192",
            temperature=0.1,
        )
    except Exception as e:
        st.error(f"❌ Error initializing LLM: {str(e)}")
        return None

def format_structured_item(item):
    """Format a structured item for display with emojis."""
    if isinstance(item, dict):
        formatted = []
        if 'type' in item:
            formatted.append(f"📍 **Type:** {item['type']}")
        
        emoji_map = {
            'findings': '🔎',
            'details': '📝',
            'duration': '⏱️',
            'notes': '📌',
            'test': '🧪',
            'region': '🎯',
            'priority': '⚡',
            'reason': '❓'
        }
        
        for key, value in item.items():
            if key != 'type':
                emoji = emoji_map.get(key, '•')
                if isinstance(value, list):
                    formatted.append(f"{emoji} **{key.title()}:**")
                    formatted.extend([f"  ▫️ {finding}" for finding in value])
                else:
                    formatted.append(f"{emoji} **{key.title()}:** {value}")
        return "\n".join(formatted)
    return str(item)

def create_prompt_template():
    """Create the prompt template for the LLM."""
    template = """Extract information from this veterinary consultation and format it as a JSON object.

Expected output format should include the following fields where available:
- 📅 date: consultation date in YYYY-MM-DD format
- 👨‍⚕️ veterinarian_name: name of the veterinarian
- 🐾 pet_name: name of the pet
- 🦮 pet_breed: breed of the pet
- ⏳ pet_age: age of the pet in years
- 👤 owner_name: name of the pet owner
- 📱 owner_phone: contact phone number
- 🤒 symptoms: array of observed symptoms
- 🔍 examinations: array of examination objects, each containing:
    - type: type of examination
    - findings: array of findings
- 💡 recommendations: array of recommendation objects, each containing:
    - type: type of recommendation
    - details: specific instructions
    - duration: time period if applicable
    - notes: additional information
- 🔬 diagnostics: array of diagnostic objects, each containing:
    - type: type of diagnostic
    - test: specific test name
    - region: body region if applicable
    - priority: urgency level
    - reason: reason for the diagnostic

Consultation text:
{text}

Please provide a well-structured JSON object containing all available information from the consultation."""
    return PromptTemplate.from_template(template)

def validate_json_response(response_content):
    """Validate and extract JSON from the LLM response."""
    try:
        json_match = re.search(r'\{[\s\S]*\}', response_content)
        if not json_match:
            raise ValueError("❌ No JSON object found in response")
        
        json_str = json_match.group(0)
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        st.error(f"❌ Invalid JSON format: {str(e)}")
        return None

def process_consultation(text, llm):
    """Process the consultation text using the LLM."""
    try:
        prompt = create_prompt_template()
        formatted_prompt = prompt.format(text=text)
        response = llm.invoke([HumanMessage(content=formatted_prompt)])
        return validate_json_response(response.content)
    except Exception as e:
        st.error(f"🚫 Error processing consultation: {str(e)}")
        return None

def create_docx_report(data):
    """Generate a DOCX report from the extracted data."""
    doc = Document()
    doc.add_heading('🏥 Veterinary Consultation Report 🐾', 0)
    
    # Add timestamp
    doc.add_paragraph(f"📅 Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph('=' * 50)
    
    sections = {
        '📌 Basic Information': ['date', 'veterinarian_name'],
        '👤 Owner Information': ['owner_name', 'owner_phone'],
        '🐾 Pet Information': ['pet_name', 'pet_breed', 'pet_age']
    }
    
    for section_title, fields in sections.items():
        doc.add_heading(section_title, level=1)
        for field in fields:
            if field in data:
                doc.add_paragraph(f"{field.replace('_', ' ').title()}: {data[field]}")
    
    if "symptoms" in data:
        doc.add_heading('🤒 Symptoms', level=1)
        for symptom in data["symptoms"]:
            doc.add_paragraph(f"⚠️ {symptom}", style='List Bullet')
    
    structured_sections = {
        '🔍 Examinations': 'examinations',
        '💡 Recommendations': 'recommendations',
        '🔬 Diagnostics': 'diagnostics'
    }
    
    for section_title, field in structured_sections.items():
        if field in data:
            doc.add_heading(section_title, level=1)
            for item in data[field]:
                doc.add_paragraph(format_structured_item(item))
    
    return doc

def create_pdf_report(data):
    """Generate a PDF report from the extracted data."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=10
    ))

    story = []
    story.append(Paragraph("🏥 Veterinary Consultation Report 🐾", styles['Title']))

    # Add timestamp
    story.append(Paragraph(f"📅 Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Add sections
    sections = {
        '📌 Basic Information': ['date', 'veterinarian_name'],
        '👤 Owner Information': ['owner_name', 'owner_phone'],
        '🐾 Pet Information': ['pet_name', 'pet_breed', 'pet_age']
    }

    for section_title, fields in sections.items():
        story.append(Paragraph(section_title, styles['CustomHeading']))
        for field in fields:
            if field in data:
                story.append(Paragraph(f"{field.replace('_', ' ').title()}: {data[field]}", styles['Normal']))
        story.append(Spacer(1, 12))

    if "symptoms" in data:
        story.append(Paragraph('🤒 Symptoms', styles['CustomHeading']))
        for symptom in data["symptoms"]:
            story.append(Paragraph(f"⚠️ {symptom}", styles['Normal']))

    structured_sections = {
        '🔍 Examinations': 'examinations',
        '💡 Recommendations': 'recommendations',
        '🔬 Diagnostics': 'diagnostics'
    }

    for section_title, field in structured_sections.items():
        if field in data:
            story.append(Paragraph(section_title, styles['CustomHeading']))
            for item in data[field]:
                story.append(Paragraph(format_structured_item(item), styles['Normal']))
            story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)  # Move to the beginning of the BytesIO buffer
    return buffer

def create_download_buttons(data):
    """Create download buttons for DOCX and PDF reports."""
    # Create DOCX
    doc = create_docx_report(data)
    docx_bio = BytesIO()
    doc.save(docx_bio)
    docx_bio.seek(0)  # Move to the beginning of the BytesIO buffer

    # Create PDF
    pdf_bio = create_pdf_report(data)

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="📄 Download DOCX",
            data=docx_bio.getvalue(),
            file_name=f"vet_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col2:
        st.download_button(
            label="📑 Download PDF",
            data=pdf_bio.getvalue(),
            file_name=f"vet_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            mime="application/pdf"
        )

def display_consultation_data(data):
    """Display the extracted consultation data in the Streamlit interface."""
    st.write("## 📊 Extracted Details")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.subheader("📌 Basic Information")
        if "date" in data:
            st.write(f"📅 Date: {data['date']}")
        if "veterinarian_name" in data:
            st.write(f"👨‍⚕️ Veterinarian: {data['veterinarian_name']}")
        
        st.subheader("👤 Owner Information")
        if "owner_name" in data:
            st.write(f"📛 Name: {data['owner_name']}")
        if "owner_phone" in data:
            st.write(f"📱 Phone: {data['owner_phone']}")

    with col2:
        st.subheader("🐾 Pet Information")
        pet_fields = {
            "pet_name": "🏷️ Name",
            "pet_breed": "🦮 Breed",
            "pet_age": "⏳ Age"
        }
        for field, label in pet_fields.items():
            if field in data:
                st.write(f"{label}: {data[field]}")
        
        if "symptoms" in data:
            st.subheader("🤒 Symptoms")
            for symptom in data["symptoms"]:
                st.write(f"⚠️ {symptom}")

    with col3:
        structured_sections = {
            "🔍 Examinations": ("examinations", "🩺"),
            "💡 Recommendations": ("recommendations", "✨"),
            "🔬 Diagnostics": ("diagnostics", "🔎")
        }
        
        for title, (field, emoji) in structured_sections.items():
            if field in data:
                st.subheader(title)
                for item in data[field]:
                    st.markdown(f"{emoji} {format_structured_item(item)}")

# Function to initialize the SQLite database
def initialize_database():
    conn = sqlite3.connect('consultations.db')  # Create a database file
    cursor = conn.cursor()
    # Create a table if it doesn't exist
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS consultations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            date TEXT,
            veterinarian_name TEXT,
            pet_name TEXT,
            pet_breed TEXT,
            pet_age INTEGER,
            owner_name TEXT,
            owner_phone TEXT,
            symptoms TEXT,
            examinations TEXT,
            recommendations TEXT,
            diagnostics TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    return conn

# Function to insert data into the database
def insert_data(conn, data):
    cursor = conn.cursor()
    # Prepare the data for insertion
    try:
        cursor.execute('''
            INSERT INTO consultations (date, veterinarian_name, pet_name, pet_breed, pet_age, owner_name, owner_phone, symptoms, examinations, recommendations, diagnostics)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            data.get('date'),
            data.get('veterinarian_name'),
            data.get('pet_name'),
            data.get('pet_breed'),
            data.get('pet_age'),
            data.get('owner_name'),
            data.get('owner_phone'),
            json.dumps(data.get('symptoms', [])),  # Convert list to JSON string
            json.dumps(data.get('examinations', [])),  # Convert list to JSON string
            json.dumps(data.get('recommendations', [])),  # Convert list to JSON string
            json.dumps(data.get('diagnostics', []))  # Convert list to JSON string
        ))
        conn.commit()
    except Exception as e:
        st.error(f"❌ Error inserting data into database: {str(e)}")

# Modify the main function to include database initialization and data insertion
def main():
    st.markdown("<h1 style='font-size: 24px;'>🏥 Veterinary Consultation Scanner 🐾</h1>", unsafe_allow_html=True)
    st.markdown("<h2 style='font-size: 18px;'>📄 Upload veterinary consultation documents to extract and analyze key information</h2>", unsafe_allow_html=True)
    
    status = st.empty()
    file = st.file_uploader("📎 Upload PDF or Word Doc", type=["pdf", "docx"])
    
    if file is not None:
        with st.spinner("🔍 Scanning document..."):
            try:
                # Extract text from document
                text = extract_text_from_file(file)
                
                # Initialize LLM
                llm = initialize_llm()
                if llm is None:
                    return  # Exit if LLM initialization failed
                
                # Process consultation
                data = process_consultation(text, llm)
                if data is None:
                    return  # Exit if processing failed
                
                # Initialize the database
                conn = initialize_database()
                
                # Insert data into the database
                insert_data(conn, data)
                
                # Create download buttons
                create_download_buttons(data)
                
                # Display results
                display_consultation_data(data)
                
                status.success("✅ Consultation Scanned Successfully")
                
            except Exception as e:
                st.error(f"❌ An error occurred: {str(e)}")
                st.exception(e)
    else:
        st.info("📤 Please upload a file to begin scanning.")

if __name__ == '__main__':
    main()